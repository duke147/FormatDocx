from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from docx import Document
import json


def is_json(myjson):
    """检查字符串是否为有效的JSON"""
    try:
        json_object = json.loads(myjson)
    except ValueError as e:
        return False

    return True


def create_table_with_json(doc, json_str):
    """在文档中创建一个新的表格，包含一个单元格，用于显示JSON字符串"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # 设置单元格的背景颜色为#c0c0c0
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(r'<w:shd {} w:fill="C0C0C0"/>'.format(nsdecls('w')))
    tcPr.append(shd)

    # 设置单元格边框为黑色
    tcBorders = parse_xml(r'''
    <w:tcBorders {}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000" />
        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000" />
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000" />
        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000" />
    </w:tcBorders>'''.format(nsdecls('w')))
    tcPr.append(tcBorders)

    # 设置字体样式
    if cell.paragraphs:  # 检查是否有段落
        paragraph = cell.paragraphs[0]
        if paragraph.runs:  # 检查段落中是否有运行对象
            font = paragraph.runs[0].font
            font.name = 'Menlo' if 'Menlo' in font.available_fonts else 'Courier New'
            font.size = Pt(10)  # 根据需要调整字体大小

    # 插入JSON字符串并保持其格式（防止中文字符转为 Unicode 编码）
    cell.text = json.dumps(json.loads(json_str), indent=4, ensure_ascii=False)

    return table


def process_document(doc_path, save_path):
    """
    同时处理段落和表格中的 JSON
    :param doc_path: 原文档路径
    :param save_path: 保存路径
    :return:
    """
    doc = Document(doc_path)

    # 首先处理文档中的段落
    process_elements(doc.paragraphs, doc)

    # 然后处理文档中的表格
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                process_elements(cell.paragraphs, doc, (i, row, cell))

    # 保存文档
    doc.save(save_path)


def process_elements(elements, doc, table_info=None):
    json_str = ""
    json_paragraphs = []  # 用于存储包含 JSON 字符串的段落
    stack = []  # 使用栈来跟踪 JSON 对象的开始和结束

    # 查找并处理未被表格包裹的JSON字符串
    for para in elements:
        for char in para.text:
            if char == "{":
                stack.append("{")  # 将 "{" 压入栈中
                if len(stack) == 1:  # 当开始新的 JSON 对象时，清空 json_str
                    json_str = "{"
                    json_paragraphs = [para]  # 当开始新的 JSON 对象时，清空 json_paragraphs
                else:
                    json_str += char
            elif char == "}":
                if stack:  # 如果栈不为空
                    stack.pop()  # 从栈中弹出一个 "{"
                    json_str += char
                    if para not in json_paragraphs:
                        json_paragraphs.append(para)  # 添加包含 JSON 字符串的段落
                    if not stack:  # 如果栈为空，表示找到一个完整的 JSON 对象
                        if is_json(json_str):
                            print(f"有效 JSON: {json_str}")
                            if not any(run._r.xml for run in para.runs if '<w:tbl' in run._r.xml):
                                # 创建新表格
                                new_table = create_table_with_json(doc, json_str)
                                if table_info:  # 如果我们正在处理表格
                                    # 替换旧表格
                                    i, row, cell = table_info
                                    old_table = doc.tables[i]
                                    p = old_table._element.getparent()
                                    p.insert(p.index(old_table._element), new_table._element)
                                    p.remove(old_table._element)
                                else:
                                    # 在第一个 json_paragraph 的位置插入新表格
                                    p = json_paragraphs[0]._element
                                    p.getparent().insert(p.getparent().index(p) + 1, new_table._element)
                                # 清除包含 JSON 字符串的段落的内容
                                for p in json_paragraphs:
                                    p.clear()
                                    p._element.getparent().remove(p._element)  # 从父元素中删除段落
                        else:
                            print(f"无效 JSON: {json_str}")
                        json_str = ""  # 重置 json_str 以处理下一段字符
                        json_paragraphs = []  # 清空 json_paragraphs 以处理下一个 JSON 对象
                else:
                    json_str += char
                    if para not in json_paragraphs:
                        json_paragraphs.append(para)  # 添加包含 JSON 字符串的段落
            elif len(stack) > 0:  # 如果栈中有 "{"
                json_str += char
                if para not in json_paragraphs:
                    json_paragraphs.append(para)  # 添加包含 JSON 字符串的段落


# 使用示例
doc_path = '/Users/wangfugui/Downloads/test.docx' # 原文档
save_path = '/Users/wangfugui/Downloads/tmp.docx' # 处理后文档
process_document(doc_path, save_path)